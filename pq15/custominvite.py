import docx
from docx.enum.text import WD_BREAK
doc = docx.Document('invitations.docx')
file = open("guests.txt")
read = file.readlines()
for name in read:
    doc.add_paragraph('It would be a pleasure to have the company of').style = 'styling'
    doc.add_paragraph(name).style = 'names'
    doc.add_paragraph('at 11010 Memory Lane on Evening of').style = 'styling'
    doc.add_paragraph('April 1st').style = 'LO-normal'
    date = doc.add_paragraph('at 7 o\'clock')
    date.runs[0].add_break(break_type=WD_BREAK.PAGE)
    date.style = 'styling'
doc.save('invitations.docx')
